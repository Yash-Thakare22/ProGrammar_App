from gingerit.gingerit import GingerIt
import flask
from flask import request
import werkzeug
from docx import Document
from docx.enum.text import WD_COLOR_INDEX

import pytesseract
import argparse
import cv2
import os
from PIL import Image

import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
import requests



app = flask.Flask(__name__)

@app.route('/', methods = ['GET', 'POST'])
def handle_request():
	imagefile = flask.request.files['image']
	filename = werkzeug.utils.secure_filename(imagefile.filename)
	print("\nReceived image File name : " + imagefile.filename)
	print(request.environ.get('HTTP_X_REAL_IP', request.remote_addr))
	
	imagefile.save(filename)
	


	#text = 'At the start ofschool Dora was afrad of her new Teacher.'
	#text = 'It the end'
	
	pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract'
	
	print("ok")
	# load the example image and convert it to grayscale
	image = cv2.imread('androidFlask.png')
	gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)

	'''
	
	# write the grayscale image to disk as a temporary file so we can
	# apply OCR to it
	filename = "{}.png".format(os.getpid())
	cv2.imwrite(filename, gray)
	'''

	# load the image as a PIL/Pillow image, apply OCR, and then delete
	# the temporary file
	text = pytesseract.image_to_string(Image.open(filename))
	

	text=text.replace('\n\n',' ')
	text=text.replace('\n',' ')
	print(text)



	parser = GingerIt()
	print(parser.parse(text)) 

	res  = parser.parse(text)['result']

	document = Document()
	document.add_paragraph('Original Text')
	document.add_paragraph(text)
	document.add_paragraph("Corrected Text")
	document.add_paragraph(res)
	document.add_paragraph("Text with errors highlighted") 

	correct =  parser.parse(text)['corrections']
	correct.reverse()
	p2 = document.add_paragraph()
	l = text.split(" ")
	for e in l:
		if len(correct)==0:
			p2.add_run(e)
			p2.add_run(' ')
		for ele in correct:
			if '.' in e:
				x = e.replace('.','')
				if x == ele['text']:
					font = p2.add_run(x).font
					font.highlight_color = WD_COLOR_INDEX.RED
					p2.add_run('.')
					p2.add_run(' ')
					correct.pop(0)
					break

				else:
					p2.add_run(x)
					p2.add_run('.')
					p2.add_run(' ')
					break
			elif ',' in e:
				x = e.replace(',','')
				if x == ele['text']:
					font = p2.add_run(x).font
					font.highlight_color = WD_COLOR_INDEX.RED
					p2.add_run(',')
					p2.add_run(' ')
					correct.pop(0)
					break

				else:
					p2.add_run(x)
					p2.add_run(',')
					p2.add_run(' ')
					break
			if '!' in e:
				x = e.replace('!','')
				if x == ele['text']:
					font = p2.add_run(x).font
					font.highlight_color = WD_COLOR_INDEX.RED
					p2.add_run('!')
					p2.add_run(' ')
					correct.pop(0)
					break

				else:
					p2.add_run(x)
					p2.add_run('!')
					p2.add_run(' ')
					break
			if '?' in e:
				x = e.replace('?','')
				if x == ele['text']:
					font = p2.add_run(x).font
					font.highlight_color = WD_COLOR_INDEX.RED
					p2.add_run('?')
					p2.add_run(' ')
					correct.pop(0)
					break

				else:
					p2.add_run(x)
					p2.add_run('?')
					p2.add_run(' ')
					break
			else:
				if e == ele['text']:
					font = p2.add_run(e).font
					font.highlight_color = WD_COLOR_INDEX.RED
					p2.add_run(' ')
					correct.pop(0)
					break

				else:
					p2.add_run(e)
					p2.add_run(' ')
					break


			


	document.save("out.docx")


	fromaddr = "programmar65@gmail.com"
	toaddr = "s.navanit@gmail.com"
	   
	# instance of MIMEMultipart 
	msg = MIMEMultipart() 
	  
	# storing the senders email address   
	msg['From'] = fromaddr 
	  
	# storing the receivers email address  
	msg['To'] = toaddr 
	  
	# storing the subject  
	msg['Subject'] = "Result from Programmar"
	  
	# string to store the body of the mail 
	body = '''
	Hello,

	Following is an attachment which contains the corrected 
	text document as per your request from our application
	Programmar

	Thank you
	'''
	  
	# attach the body with the msg instance 
	msg.attach(MIMEText(body, 'plain')) 
	  
	# open the file to be sent  
	filename = "out.docx"
	attachment = open("out.docx", "rb") 
	  
	# instance of MIMEBase and named as p 
	p = MIMEBase('application', 'octet-stream') 
	  
	# To change the payload into encoded form 
	p.set_payload((attachment).read()) 
	  
	# encode into base64 
	encoders.encode_base64(p) 
	   
	p.add_header('Content-Disposition', "attachment; filename= %s" % filename) 
	  
	# attach the instance 'p' to instance 'msg' 
	msg.attach(p) 
	  
	# creates SMTP session 
	s = smtplib.SMTP('smtp.gmail.com', 587) 
	  
	# start TLS for security 
	s.starttls() 
	  
	# Authentication 
	s.login(fromaddr, "Programmar123") 
	  
	# Converts the Multipart msg into a string 
	text = msg.as_string() 
	  
	# sending the mail 
	s.sendmail(fromaddr, toaddr, text) 
	  
	# terminating the session 
	s.quit() 
	
	return "Please Check Email"
	
if __name__ == "__main__":
	app.debug = True 
	app.run(host='0.0.0.0', port = 5000)
